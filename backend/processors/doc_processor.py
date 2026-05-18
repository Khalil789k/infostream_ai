

from docx import Document
from typing import Optional
import io


class DocProcessor:
 
    
    @staticmethod
    def extract_text(file_path: Optional[str] = None, file_bytes: Optional[bytes] = None) -> str:
       
        if not file_path and not file_bytes:
            raise ValueError("Either file_path or file_bytes must be provided")
        
        try:
            if file_bytes:
                doc = Document(io.BytesIO(file_bytes))
            else:
                doc = Document(file_path)
            
            text_content = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            full_text = '\n\n'.join(text_content)
            
            if not full_text.strip():
                raise Exception("No text content found in document")
            
            return full_text.strip()
            
        except Exception as e:
            raise Exception(f"Failed to extract text from document: {str(e)}")

